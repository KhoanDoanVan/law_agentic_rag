import tiktoken
from langchain.text_splitter import RecursiveCharacterTextSplitter
from pathlib import Path
from typing import List, Dict, Any, Optional
import os
import json
import subprocess


class LegalDocumentProcessor:
  """Xử lý và chuẩn bị documents trước khi embedding"""

  def __init__(self, chunk_size: int = 5000, chunk_overlap: int = 500):  # 12.5% overlap.
    self.chunk_size = chunk_size
    self.chunk_overlap = chunk_overlap
    self.text_splitter = RecursiveCharacterTextSplitter(
      chunk_size=chunk_size,
      chunk_overlap=chunk_overlap,
      separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ": ", " ", ""]  # Thêm space sau dấu chấm
    )
    self.encoding = tiktoken.get_encoding("cl100k_base")


  def extract_text_from_file(self, file_path: str) -> str:
    """Trích xuất text từ các loại file khác nhau"""

    file_ext = Path(file_path).suffix.lower()

    try:
      if file_ext == ".txt":
        with open(file_path, 'r', encoding='utf-8') as f:
          return f.read()

      elif file_ext == ".pdf":
        import PyPDF2
        with open(file_path, 'rb') as f:
          reader = PyPDF2.PdfReader(f)
          text = ""
          for page in reader.pages:
            text += page.extract_text() + "\n"
          return text

      elif file_ext == ".docx":
        from docx import Document
        try:
          filename = Path(file_path).name
          if filename.startswith(("~$", "._")):
            print(f"⚠️ Skip temp/hidden file: {file_path}")
            return ""
          doc = Document(file_path)
          # Sửa lại để extract text tốt hơn
          full_text = []
          for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
              full_text.append(paragraph.text.strip())
          
          # Join với double newline để preserve structure
          return "\n\n".join(full_text)
        except Exception as e:
          print(f"❌ Error reading {file_path}: {e}")
          return ""
        
      elif file_ext == ".doc":
        try:
          out_dir = str(Path(file_path).parent)
          subprocess.run(
            ["soffice", "--headless", "--convert-to", "docx", file_path, "--outdir", out_dir],
            check=True
          )
          # Tạo đường dẫn mới .docx
          converted_path = Path(file_path).with_suffix(".docx")
          if not converted_path.exists():
            print(f"❌ Không convert được {file_path}")
            return ""
          # Đọc docx mới
          from docx import Document
          doc = Document(str(converted_path))
          text = "\n".join([p.text for p in doc.paragraphs])

          os.remove(file_path)

          return text
        except FileNotFoundError:
          print("❌ LibreOffice (soffice) chưa được cài đặt. Cài bằng: brew install libreoffice hoặc apt-get install libreoffice")
          return ""

      else:
        return ""

    except Exception as e:
      print(f"Error processing {file_path}: {e}")
      return ""


  def load_folder_metadata(self, folder_path: str) -> Optional[Dict]:
    """Load metadata từ file meta.json trong folder"""

    meta_file = os.path.join(folder_path, "meta.json")
    if os.path.exists(meta_file):
      try:
        with open(meta_file, 'r', encoding='utf-8') as f:
          return json.load(f)
      except Exception as e:
        print(f"Error loading metadata from {meta_file}: {e}")

    return None


  def chunk_document(self, text: str, metadata: Dict) -> List[Dict]:
    """Chia document thành chunks với metadata đầy đủ"""

    chunks = self.text_splitter.split_text(text)

    chunk_docs = []

    for i, chunk in enumerate(chunks):
      chunk_id = f"{metadata['document_id']}_chunk_{i}"

      chunk_metadata = {
        **metadata,
        'chunk_id': chunk_id,
        'chunk_index': i,
        'total_chunks': len(chunks),
        'chunk_length': len(chunk),
        'token_count': len(self.encoding.encode(chunk))
      }

      chunk_docs.append({
        'id': chunk_id,
        'text': chunk,
        'metadata': chunk_metadata
      })

    return chunk_docs
  